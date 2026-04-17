"""
Markdown -> DOCX conversion with Altus grant application branding.

Converts the final markdown grant draft into a professionally formatted Word document.

Adapted from Engage Together POC -- same rendering engine, grant-specific title page.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# Altus brand colors
ALTUS_BLUE = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
ALTUS_TEAL = RGBColor(0x00, 0x96, 0x88)  # Accent teal
ALTUS_GRAY = RGBColor(0x6B, 0x6B, 0x6B)  # Body gray


def markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    title: str = "Grant Application",
    subtitle: str = "",
    applicant: str = "Altus Solutions",
    funder: str = "",
) -> Path:
    """
    Convert markdown grant draft to a formatted DOCX file.

    Args:
        markdown_text: The markdown grant content.
        output_path: Where to save the .docx file.
        title: Document title for the title page.
        subtitle: Opportunity number or program name.
        applicant: Applicant organization name.
        funder: Funder agency name.

    Returns:
        Path to the created DOCX file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = ALTUS_BLUE
        heading_style.font.name = "Calibri"
        if level == 1:
            heading_style.font.size = Pt(20)
            heading_style.font.bold = True
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        elif level == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.bold = True
        elif level == 4:
            heading_style.font.size = Pt(11)
            heading_style.font.bold = True

    # --- Title Page ---
    doc.add_paragraph("")
    doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = ALTUS_BLUE
    run.bold = True

    if subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = ALTUS_TEAL

    doc.add_paragraph("")

    if funder:
        funder_para = doc.add_paragraph()
        funder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = funder_para.add_run(f"Submitted to: {funder}")
        run.font.size = Pt(12)
        run.font.color.rgb = ALTUS_GRAY

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_para.add_run(f"Applicant: {applicant}")
    run.font.size = Pt(14)
    run.font.color.rgb = ALTUS_GRAY
    run.italic = True

    doc.add_page_break()

    # --- Parse and render markdown ---
    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # Skip title page elements if they duplicate what we already added
        if i < 5 and (line.startswith("# ") and title.lower() in line.lower()):
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        elif line.startswith("### "):
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith("## "):
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith("# "):
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Table rows
        if "|" in line and line.strip().startswith("|"):
            if re.match(r"^\|[\s\-:]+\|", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            _flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^\s*(\d+)\.\s+(.+)", line)
        if num_match:
            text = num_match.group(2)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraphs
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1

    # Flush any remaining table
    _flush_table(doc, table_rows)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _flush_table(doc: Document, rows: list[list[str]]) -> None:
    """Render accumulated table rows into a DOCX table."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    rows.clear()


def _add_formatted_text(para, text: str) -> None:
    """Parse inline markdown formatting (bold, italic, code) and add to paragraph."""
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"  # bold+italic
        r"|(\*\*(.+?)\*\*)"     # bold
        r"|(\*(.+?)\*)"         # italic
        r"|(`(.+?)`)"           # code
    )

    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])

        if match.group(2):  # bold+italic
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # bold
            run = para.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # italic
            run = para.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # code
            run = para.add_run(match.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)

        last_end = match.end()

    if last_end < len(text):
        para.add_run(text[last_end:])
