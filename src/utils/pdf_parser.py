"""
PDF and DOCX text extraction for RFP documents.

Uses PyMuPDF (fitz) for PDF text extraction and pdfplumber for table extraction.
Falls back gracefully when documents have complex layouts.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument


@dataclass
class ExtractedTable:
    """A table extracted from the document."""
    page_number: int
    headers: list[str]
    rows: list[list[str]]


@dataclass
class ExtractedDocument:
    """Complete extraction result from an RFP document."""
    full_text: str
    page_count: int
    pages: list[str]  # text per page
    tables: list[ExtractedTable] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


def extract_pdf(file_path: Path) -> ExtractedDocument:
    """
    Extract text and tables from a PDF file.

    Uses PyMuPDF for text (fast, layout-aware) and pdfplumber for tables.
    """
    pages: list[str] = []
    metadata: dict = {}

    # --- Text extraction with PyMuPDF ---
    doc = fitz.open(str(file_path))
    metadata = {
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "subject": doc.metadata.get("subject", ""),
        "creator": doc.metadata.get("creator", ""),
        "page_count": len(doc),
    }

    for page in doc:
        text = page.get_text("text")
        # Clean up common PDF artifacts
        text = re.sub(r"\n{3,}", "\n\n", text)  # collapse excessive newlines
        text = re.sub(r"[ \t]{2,}", " ", text)   # collapse spaces
        pages.append(text.strip())
    doc.close()

    full_text = "\n\n--- Page Break ---\n\n".join(
        f"[Page {i+1}]\n{page}" for i, page in enumerate(pages)
    )

    # --- Table extraction with pdfplumber ---
    tables: list[ExtractedTable] = []
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for raw_table in page_tables:
                    if not raw_table or len(raw_table) < 2:
                        continue
                    # First row as headers
                    headers = [str(c or "").strip() for c in raw_table[0]]
                    rows = [
                        [str(c or "").strip() for c in row]
                        for row in raw_table[1:]
                    ]
                    tables.append(ExtractedTable(
                        page_number=i + 1,
                        headers=headers,
                        rows=rows,
                    ))
    except Exception:
        # pdfplumber can fail on some PDFs -- text extraction still works
        pass

    return ExtractedDocument(
        full_text=full_text,
        page_count=len(pages),
        pages=pages,
        tables=tables,
        metadata=metadata,
    )


def extract_docx(file_path: Path) -> ExtractedDocument:
    """Extract text and tables from a DOCX file."""
    doc = DocxDocument(str(file_path))

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    level_int = int(level)
                    text = "#" * level_int + " " + text
                except ValueError:
                    pass
            paragraphs.append(text)

    full_text = "\n\n".join(paragraphs)

    # Extract tables
    tables: list[ExtractedTable] = []
    for i, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]
        tables.append(ExtractedTable(
            page_number=0,  # DOCX doesn't have page numbers easily
            headers=headers,
            rows=rows,
        ))

    return ExtractedDocument(
        full_text=full_text,
        page_count=1,  # Approximation
        pages=[full_text],
        tables=tables,
        metadata={"title": doc.core_properties.title or ""},
    )


def extract_document(file_path: Path) -> ExtractedDocument:
    """Extract text from a PDF or DOCX file, auto-detecting format."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Use PDF or DOCX.")


def format_tables_as_markdown(tables: list[ExtractedTable]) -> str:
    """Convert extracted tables to markdown format for LLM consumption."""
    if not tables:
        return ""

    parts: list[str] = []
    for i, table in enumerate(tables):
        parts.append(f"\n### Table {i+1} (Page {table.page_number})\n")
        if table.headers:
            parts.append("| " + " | ".join(table.headers) + " |")
            parts.append("| " + " | ".join("---" for _ in table.headers) + " |")
        for row in table.rows:
            parts.append("| " + " | ".join(row) + " |")

    return "\n".join(parts)
